"""
generate_docs.py — Creates the Hyperlocal Network project documentation as a .docx file.
Run: python generate_docs.py
Output: Hyperlocal_Network_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x2e)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    # light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table_with_headers(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, '4F46E5')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size  = Pt(10)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8F8FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Document ────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('HYPERLOCAL NETWORK')
title_run.font.size  = Pt(32)
title_run.font.bold  = True
title_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_para.add_run('Partnership Platform')
sub_run.font.size  = Pt(20)
sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()

tagline_para = doc.add_paragraph()
tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_run = tagline_para.add_run('Product Documentation, Flow Guide & Implementation Manual')
tag_run.font.size  = Pt(13)
tag_run.font.italic = True
tag_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Version 1.0  •  {datetime.date.today().strftime("%B %Y")}')
date_run.font.size  = Pt(11)
date_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (manual, since python-docx TOC field needs Word to update)
# ════════════════════════════════════════════════════════

add_heading(doc, 'Table of Contents', level=1, color=(0x4F, 0x46, 0xE5))

toc_entries = [
    ('1', 'What Is Hyperlocal Network?'),
    ('2', 'System Architecture'),
    ('3', 'Complete Product Flowchart'),
    ('4', 'Feature Deep-Dive'),
    ('4.1', '  Brand Registration & Login'),
    ('4.2', '  Find Partners (Discovery)'),
    ('4.3', '  Create a Partnership Proposal'),
    ('4.4', '  Partnership Lifecycle & Statuses'),
    ('4.5', '  Different Offer Mode'),
    ('4.6', '  Terms & Rules Negotiation'),
    ('4.7', '  Partnership Alerts (Notifications)'),
    ('4.8', '  Redeem Token (Point Issuance)'),
    ('4.9', '  Campaigns & Follow-ups'),
    ('4.10', '  Hyperlocal Networks (Groups)'),
    ('4.11', '  Partner Offers'),
    ('4.12', '  Growth & Analytics'),
    ('4.13', '  Customer Portal'),
    ('5', 'Database Schema Summary'),
    ('6', 'API Reference'),
    ('7', 'Implementation Guide (New PC Setup)'),
    ('8', 'Glossary'),
]

for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 if '.' in num else 0)
    r = p.add_run(f'{num}   {title}')
    r.font.size = Pt(10.5)
    if '.' not in num:
        r.font.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS HYPERLOCAL NETWORK?
# ════════════════════════════════════════════════════════

add_heading(doc, '1. What Is Hyperlocal Network?', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_para(doc,
    'Hyperlocal Network is a B2B SaaS platform that allows local merchant brands — '
    'cafes, gyms, salons, restaurants, pharmacies and more — to form "partnerships" with '
    'each other. Through a partnership, Brand A can reward Brand B\'s customers and vice versa, '
    'increasing footfall for both.',
    size=11)

doc.add_paragraph()
add_para(doc, 'The Big Idea (Plain English)', bold=True, size=12, color=(0x4F, 0x46, 0xE5))
add_para(doc,
    'Imagine you run a gym. Your members also go to nearby cafes. Instead of running expensive ads, '
    'you partner with that cafe: when a customer spends ₹500 at the cafe they earn 50 points at YOUR '
    'gym — bringing them through your door. Both businesses win. This is what Hyperlocal Network enables.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Core Value Propositions', bold=True, size=12, color=(0x4F, 0x46, 0xE5))
add_bullet(doc, 'Zero ad spend customer acquisition via partner brands')
add_bullet(doc, 'Automated loyalty point exchange between participating outlets')
add_bullet(doc, 'Full control over discount type, caps, and minimum spends')
add_bullet(doc, 'Real-time analytics on partner-sourced revenue')
add_bullet(doc, 'Network groups for multi-brand ecosystems (malls, food courts, etc.)')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════

add_heading(doc, '2. System Architecture', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_para(doc, 'Technology Stack', bold=True, size=12, color=(0x4F, 0x46, 0xE5))

add_table_with_headers(doc,
    ['Layer', 'Technology', 'Port / Path', 'Purpose'],
    [
        ['Frontend', 'Vue 3 + Vite + TypeScript + Tailwind CSS', 'localhost:3000', 'Merchant dashboard SPA'],
        ['Backend API', 'Laravel 13 (PHP 8.2)', 'localhost:8000', 'RESTful API, business logic'],
        ['Database', 'SQLite (dev) / MySQL (prod)', 'database.sqlite', 'All persistent data'],
        ['Auth', 'Laravel Sanctum', '/api/auth/*', 'Token-based auth for merchant + super-admin'],
        ['State Mgmt', 'Pinia', 'Browser', 'Frontend reactive state stores'],
        ['HTTP Client', 'Axios', 'src/services/api.ts', 'Frontend API communication'],
    ],
    col_widths=[1.2, 2.2, 1.6, 2.5]
)

doc.add_paragraph()
add_para(doc, 'Module Architecture (Backend)', bold=True, size=12, color=(0x4F, 0x46, 0xE5))
add_para(doc,
    'The backend follows a Domain-Driven Design using Laravel modules. Each feature lives in its own '
    'module under app/Modules/ with Controllers, Services, Models, and Resources:',
    size=11)

add_code_block(doc, '''src/backend/app/Modules/
├── Admin/            ← Login, registration, user management
├── Partnership/      ← Core partnership lifecycle
├── Discovery/        ← Find partner suggestions & search
├── Network/          ← Hyperlocal groups (networks of brands)
├── Campaign/         ← WhatsApp/notification campaigns
├── Analytics/        ← ROI, attribution, retention stats
├── Growth/           ← Health scores, referral links, profiles
├── Customer/         ← Customer data upload & lookup
├── PartnerOffers/    ← Sharable discount offers
├── Ledger/           ← Credit/debit transaction ledger
├── Execution/        ← Token redemption at POS
├── Enablement/       ← Outlet training completion
├── EventTriggers/    ← Webhook event-based automations
└── SuperAdmin/       ← Platform admin controls''')

doc.add_paragraph()
add_para(doc, 'Frontend Module Architecture', bold=True, size=12, color=(0x4F, 0x46, 0xE5))

add_code_block(doc, '''src/frontend/src/
├── modules/
│   ├── auth/             ← Login & registration views
│   ├── partnership/      ← Partnership list + detail views
│   ├── discovery/        ← Find Partners view
│   ├── network/          ← Network views + invite landing page
│   ├── campaigns/        ← Campaign creation & list
│   ├── analytics/        ← Charts, ROI views
│   ├── growth/           ← Growth dashboard
│   ├── admin/            ← Super-admin login
│   └── ...
├── stores/               ← Pinia stores (auth, partnership, etc.)
├── services/api.ts       ← Axios HTTP client (auto-attaches token)
├── router/index.ts       ← Vue Router with auth guards
└── components/
    └── AppLayout.vue     ← Main shell: sidebar + notification bell''')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 3 — COMPLETE PRODUCT FLOWCHART
# ════════════════════════════════════════════════════════

add_heading(doc, '3. Complete Product Flowchart', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_para(doc,
    'The diagram below shows every step a merchant takes — from signing up to running a live partnership '
    'with a partner brand. Each box is a screen or system action.',
    size=11)

doc.add_paragraph()

# FLOWCHART as a formatted text table
add_para(doc, '── ONBOARDING ──────────────────────────────────────────────────────', bold=True, size=10, color=(0x4F, 0x46, 0xE5))
flow_onboarding = [
    ['[START]', '→', 'Brand visits platform'],
    ['', '↓', ''],
    ['Register Brand', '→', 'Enter: name, email, password, city, category'],
    ['', '↓', ''],
    ['Login', '→', 'Enter credentials → receive Bearer token'],
    ['', '↓', ''],
    ['Dashboard', '→', 'See overview: partnerships, recent activity'],
]
for step in flow_onboarding:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'  {step[0]:25}  {step[1]}  {step[2]}')
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)

doc.add_paragraph()
add_para(doc, '── FINDING A PARTNER ───────────────────────────────────────────────', bold=True, size=10, color=(0x4F, 0x46, 0xE5))
flow_find = [
    ['Find Partners page', '→', 'Enter city + optional category → Search'],
    ['', '↓', ''],
    ['Search Results', '→', 'List of brands with fit score & trust score'],
    ['', '↓', ''],
    ['Already partnered?', '→', 'YES → "View Partnership" button shown (green)'],
    ['', '↓', 'NO → "Propose tie-up" button shown'],
    ['', '↓', ''],
    ['Propose tie-up', '→', 'Opens create-partnership modal pre-filled with brand'],
]
for step in flow_find:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'  {step[0]:25}  {step[1]}  {step[2]}')
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)

doc.add_paragraph()
add_para(doc, '── CREATING A PARTNERSHIP ──────────────────────────────────────────', bold=True, size=10, color=(0x4F, 0x46, 0xE5))
flow_create = [
    ['Create Partnership', '→', 'Set: name, scope (brand/outlet), offer structure'],
    ['', '↓', ''],
    ['Offer Structure?', '→', '"Same" → Both brands use one shared offer config'],
    ['', '↓', '"Different" → Each brand sets their own offer'],
    ['', '↓', ''],
    ['Set Offer', '→', 'Type: Flat ₹ or Percentage (%), max cap, min bill'],
    ['', '↓', ''],
    ['Add Terms', '→', 'Optional: per-bill cap %, monthly cap, approval mode'],
    ['', '↓', ''],
    ['Submit', '→', 'Status → REQUESTED, partner gets in-app alert'],
]
for step in flow_create:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'  {step[0]:25}  {step[1]}  {step[2]}')
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)

doc.add_paragraph()
add_para(doc, '── PARTNERSHIP ACCEPTANCE FLOW ─────────────────────────────────────', bold=True, size=10, color=(0x4F, 0x46, 0xE5))
flow_accept = [
    ['Partner receives alert', '→', 'Sees amber banner: "Review this request"'],
    ['', '↓', ''],
    ['Offer Structure = Different?', '→', 'YES → Fill own offer form (type, %, caps)'],
    ['', '↓', '         System alert sent to proposer'],
    ['', '↓', 'NO → View shared offer details'],
    ['', '↓', ''],
    ['Happy with terms?', '→', 'YES → Accept → Status → AGREED'],
    ['', '↓', 'NO → Edit terms → Status → NEGOTIATING'],
    ['', '↓', '         Other party gets "terms updated" alert'],
    ['', '↓', ''],
    ['Both agree', '→', 'Go Live → Status → LIVE'],
    ['', '↓', ''],
    ['[PARTNERSHIP ACTIVE]', '→', 'Customers can now earn/redeem across brands'],
]
for step in flow_accept:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'  {step[0]:28}  {step[1]}  {step[2]}')
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)

doc.add_paragraph()
add_para(doc, '── LIVE OPERATIONS ─────────────────────────────────────────────────', bold=True, size=10, color=(0x4F, 0x46, 0xE5))
flow_live = [
    ['Customer visits Brand A', '→', 'Cashier opens Redeem Token screen'],
    ['', '↓', ''],
    ['Enter token / phone', '→', 'System looks up active partnerships'],
    ['', '↓', ''],
    ['Select Partnership', '→', 'Enter bill amount'],
    ['', '↓', ''],
    ['Calculate reward', '→', 'Apply discount rules (%, flat, caps)'],
    ['', '↓', ''],
    ['Confirm redemption', '→', 'Points credited / discount applied'],
    ['', '↓', ''],
    ['Ledger updated', '→', 'Both merchants see transaction in ledger'],
]
for step in flow_live:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'  {step[0]:28}  {step[1]}  {step[2]}')
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 4 — FEATURE DEEP-DIVE
# ════════════════════════════════════════════════════════

add_heading(doc, '4. Feature Deep-Dive', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

# ── 4.1 Auth ────────────────────────────────────────────

add_heading(doc, '4.1  Brand Registration & Login', level=2)

add_para(doc, 'Registration Flow', bold=True, size=11)
add_bullet(doc, 'Merchant provides: name, email, password, city, business category')
add_bullet(doc, 'System creates Merchant + admin User records in SQLite')
add_bullet(doc, 'Returns Sanctum Bearer token — stored in browser (Pinia auth store)')
add_bullet(doc, 'Self-registration queue: brands can register publicly; Super Admin reviews and approves')

doc.add_paragraph()
add_para(doc, 'Login Flow', bold=True, size=11)
add_bullet(doc, 'POST /api/auth/login with email + password')
add_bullet(doc, 'Returns token + user + merchant profile')
add_bullet(doc, 'Token saved in localStorage via Pinia; Axios attaches it as Authorization: Bearer <token>')
add_bullet(doc, 'Route guard in Vue Router checks auth.token before every protected page')

doc.add_paragraph()
add_para(doc, 'Roles', bold=True, size=11)
add_table_with_headers(doc,
    ['Role', 'Code', 'Can Do'],
    [
        ['Admin', '1', 'Everything: create partnerships, manage settings, run campaigns'],
        ['Manager', '2', 'View and manage partnerships; cannot change billing settings'],
        ['Cashier', '3', 'Redeem tokens only'],
        ['Super Admin', 'SA', 'Platform-wide: approve brands, allocate WhatsApp credits, view all merchants'],
    ],
    col_widths=[1.2, 0.8, 4.5]
)

doc.add_paragraph()

# ── 4.2 Find Partners ───────────────────────────────────

add_heading(doc, '4.2  Find Partners (Discovery)', level=2)

add_para(doc,
    'The Discovery module helps merchants find brands that would make good partners based on '
    'geographic proximity (same city) and category complementarity.',
    size=11)

doc.add_paragraph()
add_para(doc, 'How Search Works', bold=True, size=11)
add_bullet(doc, 'Input: city name + optional category filter')
add_bullet(doc, 'Backend: queries all active merchants in that city')
add_bullet(doc, 'Fit scoring: complementary category pairs score higher (e.g., gym + smoothie = Strong Fit)')
add_bullet(doc, 'Trust score: based on existing reviews from partner merchants')

doc.add_paragraph()
add_para(doc, 'Already-Partnered Brands', bold=True, size=11)
add_bullet(doc, 'Brands you are ALREADY partnered with still appear in results')
add_bullet(doc, 'They show a green badge ("Live partner" / "Proposed partner" / etc.)')
add_bullet(doc, '"View Partnership" button navigates directly to the partnership detail page')
add_bullet(doc, 'New brands show "Propose tie-up" button')

doc.add_paragraph()
add_para(doc, 'Fit Tiers', bold=True, size=11)
add_table_with_headers(doc,
    ['Tier', 'Label', 'Score Range', 'Badge Colour'],
    [
        ['1', 'Strong Fit', '75–100', 'Green'],
        ['2', 'Good Fit', '50–74', 'Blue'],
        ['3', 'Possible Fit', '25–49', 'Grey'],
        ['–', 'Unscored', '0–24', 'Light Grey'],
    ],
    col_widths=[0.6, 1.2, 1.5, 3.2]
)

doc.add_paragraph()

# ── 4.3 Create Partnership ──────────────────────────────

add_heading(doc, '4.3  Create a Partnership Proposal', level=2)

add_para(doc, 'Fields & What They Mean', bold=True, size=11)
add_table_with_headers(doc,
    ['Field', 'Required', 'Meaning'],
    [
        ['Partnership Name', 'Yes', 'Friendly label (e.g., "Cafe + Gym Cross-Promo")'],
        ['Partner Brand', 'Yes', 'The other merchant to invite'],
        ['Scope Type', 'Yes', 'Brand-wide (all outlets) or Outlet-level (specific branches)'],
        ['Offer Structure', 'Yes', 'Same = one shared offer; Different = each brand sets own offer'],
        ['Discount Type', 'Yes', 'Flat ₹ amount OR Percentage (%) of bill'],
        ['Discount Value', 'Yes', 'The ₹ or % value to give partner customers'],
        ['Max Cap per Bill', 'Conditional', 'Only for Percentage mode — maximum ₹ discount per transaction'],
        ['Min Bill Amount', 'No', 'Customer must spend at least this much to qualify'],
        ['Monthly Cap', 'No', 'Max total ₹ this merchant will give out per month'],
        ['Start Date / End Date', 'No', 'Optional date range for the partnership'],
        ['Approval Mode', 'No', 'Auto = cashier can redeem without approval; Manual = manager must approve each'],
        ['Terms', 'No', 'Per-bill cap %, monthly cap, and approval mode as negotiable terms'],
    ],
    col_widths=[1.8, 1.0, 3.7]
)

doc.add_paragraph()

# ── 4.4 Partnership Lifecycle ───────────────────────────

add_heading(doc, '4.4  Partnership Lifecycle & Statuses', level=2)

add_para(doc,
    'Every partnership moves through a defined set of statuses. Only certain transitions are allowed.',
    size=11)

doc.add_paragraph()
add_table_with_headers(doc,
    ['Status', 'Code', 'Meaning', 'Who Can Transition'],
    [
        ['Requested', '2', 'Proposer has sent the invitation; awaiting acceptor review', 'Acceptor can Accept / Reject; Proposer can update terms'],
        ['Negotiating', '3', 'Either party edited terms — counter-offer in progress', 'Either party can Accept or continue negotiating'],
        ['Agreed', '4', 'Both parties accepted; not yet live', 'Either party can trigger "Go Live"'],
        ['Live', '5', 'Active — customers can earn & redeem', 'Either party can Pause'],
        ['Paused', '6', 'Temporarily halted', 'Either party can Resume'],
        ['Rejected', 'R', 'Acceptor declined the proposal (terminal)', '—'],
        ['Expired', 'E', 'End date passed (terminal)', 'System-set automatically'],
    ],
    col_widths=[1.2, 0.6, 2.8, 1.9]
)

doc.add_paragraph()
add_para(doc, 'Status Transition Diagram', bold=True, size=11, color=(0x4F, 0x46, 0xE5))
add_code_block(doc, '''REQUESTED ──(accept)──► AGREED ──(go-live)──► LIVE ──(pause)──► PAUSED
    │                                                              │
    │ (terms edited                                                │ (resume)
    │  by either)                                                  ↓
    ↓                                                             LIVE
NEGOTIATING ──(both agree)──► AGREED
    │
    ├──(reject)──► REJECTED
    └──(expire)──► EXPIRED''')

doc.add_paragraph()

# ── 4.5 Different Offer Mode ────────────────────────────

add_heading(doc, '4.5  Different Offer Mode', level=2)

add_para(doc,
    'When "offer_structure = different" is selected, each participating brand independently configures '
    'the discount they will offer to the OTHER brand\'s customers. This allows asymmetric deals.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Example:', bold=True, size=11)
add_bullet(doc, 'Brand A (Cafe): offers 10% off to gym members')
add_bullet(doc, 'Brand B (Gym): offers ₹200 flat off to cafe customers')
add_bullet(doc, 'Each brand sets their own terms independently')

doc.add_paragraph()
add_para(doc, 'Acceptor Offer Fill Flow', bold=True, size=11)
add_numbered(doc, 'Proposer creates partnership with offer_structure = "different" and fills their own offer')
add_numbered(doc, 'Acceptor sees an amber banner: "This partnership uses different offers — please fill your offer"')
add_numbered(doc, 'Acceptor fills their offer inline (type, amount/%, max cap, min bill, monthly cap)')
add_numbered(doc, 'System sends in-app alert to Proposer: "[Brand] has filled their offer"')
add_numbered(doc, 'Proposer reviews and can proceed to Accept/Go Live')

doc.add_paragraph()

# ── 4.6 Terms & Negotiation ─────────────────────────────

add_heading(doc, '4.6  Terms & Rules Negotiation', level=2)

add_para(doc,
    'Both parties can edit partnership terms at any point while the status is Requested or Negotiating. '
    'Editing terms automatically moves the partnership to "Negotiating" and notifies the other party.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Editable Terms Fields', bold=True, size=11)
add_table_with_headers(doc,
    ['Term', 'Type', 'Description'],
    [
        ['Per-bill cap %', 'Percentage', 'Maximum percentage of the bill that can be discounted'],
        ['Per-bill cap ₹', 'Amount', 'Maximum ₹ discount per single transaction'],
        ['Monthly cap ₹', 'Amount', 'Maximum total ₹ discount this merchant gives out per month'],
        ['Min bill ₹', 'Amount', 'Minimum bill amount for a discount to apply'],
        ['Approval mode', 'Toggle', '1 = auto-approve redemptions; 2 = manual manager approval'],
    ],
    col_widths=[1.8, 1.2, 3.5]
)

doc.add_paragraph()
add_para(doc, 'Note:', bold=True, size=11)
add_para(doc,
    'Points-based input has been intentionally removed. All caps and minimums are entered in ₹ only.',
    size=11, italic=True)

doc.add_paragraph()

# ── 4.7 Alerts ──────────────────────────────────────────

add_heading(doc, '4.7  Partnership Alerts (Notifications)', level=2)

add_para(doc,
    'The notification bell in the sidebar shows in-app alerts for partnership activity. '
    'The system polls every 60 seconds for new alerts.',
    size=11)

doc.add_paragraph()
add_table_with_headers(doc,
    ['Alert Type', 'When Triggered', 'Recipient'],
    [
        ['offer_filled', 'Acceptor fills their offer (different mode)', 'Proposer'],
        ['offer_updated', 'Acceptor updates their already-filled offer', 'Proposer'],
        ['terms_updated', 'Either party edits terms → status moves to Negotiating', 'The OTHER party'],
        ['partner_accepted', 'Acceptor accepts the partnership', 'Proposer'],
        ['partner_rejected', 'Acceptor rejects the partnership', 'Proposer'],
    ],
    col_widths=[1.8, 2.8, 2.0]
)

doc.add_paragraph()
add_para(doc, 'Notification Bell Features', bold=True, size=11)
add_bullet(doc, 'Red badge shows unread count (shows "9+" if more than 9 unread)')
add_bullet(doc, 'Click alert → marks as read + navigates to that partnership')
add_bullet(doc, '"Mark all read" button clears all unread alerts at once')
add_bullet(doc, 'Dropdown closes when clicking anywhere outside it')
add_bullet(doc, 'Interval cleared on component unmount — no memory leaks')

doc.add_paragraph()

# ── 4.8 Redeem Token ────────────────────────────────────

add_heading(doc, '4.8  Redeem Token (Point Issuance)', level=2)

add_para(doc,
    'The Redeem Token screen is where a cashier at Brand A gives a discount/points to a customer '
    'who is a member of partner Brand B.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Steps', bold=True, size=11)
add_numbered(doc, 'Cashier opens the Redeem screen')
add_numbered(doc, 'Enters customer phone number or loyalty token')
add_numbered(doc, 'System finds active partnerships and eligibility')
add_numbered(doc, 'Cashier selects partnership and enters bill amount')
add_numbered(doc, 'System calculates discount (respecting min-bill, caps)')
add_numbered(doc, 'Cashier confirms → transaction logged in Ledger')

doc.add_paragraph()
add_para(doc, 'Approval Modes', bold=True, size=11)
add_bullet(doc, 'Auto: discount applied immediately on cashier confirm')
add_bullet(doc, 'Manual: cashier submits an approval request; manager approves/rejects from dashboard')

doc.add_paragraph()

# ── 4.9 Campaigns ───────────────────────────────────────

add_heading(doc, '4.9  Campaigns & Follow-ups', level=2)

add_para(doc,
    'Merchants can send WhatsApp campaigns to their customer base or to specific segments.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Campaign Types', bold=True, size=11)
add_bullet(doc, 'One-shot broadcast — send to all / filtered customers at a scheduled time')
add_bullet(doc, 'Follow-up campaigns — trigger-based (e.g., "customer hasn\'t visited in 30 days")')

doc.add_paragraph()
add_para(doc, 'Campaign Status Flow', bold=True, size=11)
add_code_block(doc, 'DRAFT → SCHEDULED → RUNNING → COMPLETED / CANCELLED')

doc.add_paragraph()
add_para(doc, 'WhatsApp Credits', bold=True, size=11)
add_bullet(doc, 'Each message send costs credits')
add_bullet(doc, 'Super Admin allocates credits to each merchant')
add_bullet(doc, 'Merchant can see their current balance before sending')
add_bullet(doc, 'Credit enforcement is configurable (WHATSAPP_CREDIT_ENFORCEMENT env key)')

doc.add_paragraph()

# ── 4.10 Networks ────────────────────────────────────────

add_heading(doc, '4.10  Hyperlocal Networks (Groups)', level=2)

add_para(doc,
    'A Network is a named group of merchants (e.g., "Forum Mall Brands", "Koramangala Food Court"). '
    'Network members can share offers and reach each other\'s customers.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Network Invitation Flow', bold=True, size=11)
add_numbered(doc, 'Network owner generates an invite link (64-char token)')
add_numbered(doc, 'Shares link with another brand (via WhatsApp, email, etc.)')
add_numbered(doc, 'Recipient opens link — shown a PUBLIC landing page (no login required)')
add_numbered(doc, 'Landing page shows: network name, description, inviting brand, member count, benefits')
add_numbered(doc, 'If not logged in: "Login to join" and "Register your brand" buttons')
add_numbered(doc, 'If logged in: "Join [Network Name]" button')
add_numbered(doc, 'On join: member added; network page shown')

doc.add_paragraph()

# ── 4.11 Partner Offers ─────────────────────────────────

add_heading(doc, '4.11  Partner Offers', level=2)

add_para(doc,
    'Partner Offers are standalone discount/reward offers a brand creates and can attach to partnerships '
    'or publish to networks. They appear on digital bills.',
    size=11)

doc.add_paragraph()
add_bullet(doc, 'Create an offer once; attach to multiple partnerships')
add_bullet(doc, 'Publish to a network — visible to all member brands\' customers')
add_bullet(doc, 'Toggle active/inactive without deleting')
add_bullet(doc, 'View impressions and claims analytics')

doc.add_paragraph()

# ── 4.12 Growth & Analytics ─────────────────────────────

add_heading(doc, '4.12  Growth & Analytics', level=2)

add_para(doc, 'Analytics Module', bold=True, size=11)
add_bullet(doc, 'Summary dashboard: total transactions, revenue, partner-sourced customers')
add_bullet(doc, 'Per-partnership ROI: transactions, revenue, retention rate')
add_bullet(doc, 'Trend charts: weekly/monthly activity over time')
add_bullet(doc, 'Attribution: tracks which redemption came via which partnership')

doc.add_paragraph()
add_para(doc, 'Growth Module', bold=True, size=11)
add_bullet(doc, 'Health scores: partnership engagement levels (Active / Declining / Dormant)')
add_bullet(doc, 'Referral links: shareable links to bring new customers via a partnership')
add_bullet(doc, 'Demand index: category-level demand signals in the merchant\'s city')
add_bullet(doc, 'Seasonal templates: pre-built campaign ideas for upcoming holidays')
add_bullet(doc, 'Weekly digest: auto-generated performance summary')

doc.add_paragraph()

# ── 4.13 Customer Portal ────────────────────────────────

add_heading(doc, '4.13  Customer Portal', level=2)

add_para(doc,
    'Customers can check their own reward balances and activity via a lightweight OTP-based portal.',
    size=11)

doc.add_paragraph()
add_bullet(doc, 'Customer enters phone number → receives OTP')
add_bullet(doc, 'Sees: total points, partner rewards available, recent transaction history')
add_bullet(doc, 'No account creation needed — phone number = identity')
add_bullet(doc, 'Rate limited to prevent OTP abuse')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 5 — DATABASE SCHEMA
# ════════════════════════════════════════════════════════

add_heading(doc, '5. Database Schema Summary', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_para(doc,
    'The system uses SQLite for local development. The database file is committed at '
    'src/backend/database/database.sqlite and contains all demo data.',
    size=11)

doc.add_paragraph()
add_para(doc, 'Core Tables', bold=True, size=12, color=(0x4F, 0x46, 0xE5))

add_table_with_headers(doc,
    ['Table', 'Key Columns', 'Purpose'],
    [
        ['merchants', 'id, name, city, category, is_active', 'One row per brand'],
        ['users', 'id, merchant_id, email, role (1/2/3)', 'Merchant staff accounts'],
        ['outlets', 'id, merchant_id, name, city, is_active', 'Physical branches'],
        ['partnerships', 'id, uuid, merchant_id, name, scope_type, offer_structure, status', 'Core partnership record'],
        ['partnership_participants', 'partnership_id, merchant_id, outlet_id, role, approval_status, offer_*', 'Which merchants/outlets are in a partnership and their offer config'],
        ['partnership_terms', 'partnership_id, per_bill_cap_percent, per_bill_cap_amount, monthly_cap_amount, min_bill_amount, approval_mode', 'Negotiated terms'],
        ['partnership_rules', 'partnership_id, ...', 'Business rules for a partnership'],
        ['partnership_alerts', 'partnership_id, recipient_merchant_id, type, title, body, read_at', 'In-app notifications'],
        ['redemptions', 'partnership_id, merchant_id, bill_amount, discount_given, customer_phone', 'Every transaction record'],
        ['networks', 'id, uuid, name, owner_merchant_id', 'Network groups'],
        ['network_members', 'network_id, merchant_id, status', 'Membership records'],
        ['network_invitations', 'token, network_id, expires_at', 'Invite links'],
        ['campaigns', 'uuid, merchant_id, status, scheduled_at', 'Broadcast campaigns'],
        ['partner_offers', 'uuid, merchant_id, title, discount_type, value', 'Sharable offers'],
        ['ledger_entries', 'merchant_id, partnership_id, type, amount', 'Debit/credit ledger'],
    ],
    col_widths=[2.0, 2.8, 2.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 6 — API REFERENCE
# ════════════════════════════════════════════════════════

add_heading(doc, '6. API Reference', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_para(doc, 'Base URL: http://localhost:8000/api', bold=True, size=11)
add_para(doc, 'Auth: All protected endpoints require header:  Authorization: Bearer <token>', size=11, italic=True)

doc.add_paragraph()
add_para(doc, 'Key Endpoints', bold=True, size=12, color=(0x4F, 0x46, 0xE5))

add_table_with_headers(doc,
    ['Method', 'Endpoint', 'Purpose'],
    [
        ['POST', '/auth/register', 'Register a new brand'],
        ['POST', '/auth/login', 'Login; returns token'],
        ['GET', '/auth/me', 'Get logged-in user + merchant'],
        ['GET', '/discovery/search?city=&category=', 'Search for partner brands'],
        ['GET', '/partnerships', 'List all partnerships for my brand'],
        ['POST', '/partnerships', 'Create a new partnership proposal'],
        ['GET', '/partnerships/{uuid}', 'Get full partnership detail'],
        ['PUT', '/partnerships/{uuid}', 'Update name/dates/terms/rules'],
        ['POST', '/partnerships/{uuid}/accept', 'Accept a partnership proposal'],
        ['POST', '/partnerships/{uuid}/go-live', 'Activate an agreed partnership'],
        ['POST', '/partnerships/{uuid}/reject', 'Reject a proposal'],
        ['POST', '/partnerships/{uuid}/pause', 'Pause a live partnership'],
        ['POST', '/partnerships/{uuid}/resume', 'Resume a paused partnership'],
        ['POST', '/partnerships/{uuid}/fill-offer', 'Acceptor fills their offer (different mode)'],
        ['GET', '/partnership-alerts', 'Get my in-app alerts'],
        ['POST', '/partnership-alerts/{id}/read', 'Mark alert as read'],
        ['POST', '/partnership-alerts/read-all', 'Mark all alerts as read'],
        ['GET', '/merchant/networks', 'List my networks'],
        ['POST', '/merchant/networks', 'Create a network'],
        ['POST', '/merchant/networks/{uuid}/invite', 'Generate invite link'],
        ['POST', '/merchant/networks/join/{token}', 'Join via invite link'],
        ['GET', '/public/network-invite/{token}', 'Preview invite (no auth)'],
        ['GET', '/analytics/summary', 'My analytics overview'],
        ['GET', '/analytics/partnerships/{uuid}', 'Per-partnership analytics'],
        ['GET', '/growth/health', 'Partnership health scores'],
        ['GET', '/campaigns', 'List campaigns'],
        ['POST', '/campaigns', 'Create campaign'],
        ['POST', '/campaigns/{uuid}/run', 'Send a campaign immediately'],
    ],
    col_widths=[0.8, 2.8, 2.7]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 7 — IMPLEMENTATION GUIDE
# ════════════════════════════════════════════════════════

add_heading(doc, '7. Implementation Guide (New PC Setup)', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_para(doc, 'Prerequisites', bold=True, size=12, color=(0x4F, 0x46, 0xE5))
add_table_with_headers(doc,
    ['Tool', 'Min Version', 'Download'],
    [
        ['PHP', '8.2', 'https://php.net/downloads'],
        ['Composer', '2.x', 'https://getcomposer.org'],
        ['Node.js', '20', 'https://nodejs.org'],
        ['Git', 'any', 'https://git-scm.com'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

doc.add_paragraph()
add_para(doc, 'Step 1 — Clone the repository', bold=True, size=11, color=(0x4F, 0x46, 0xE5))
add_code_block(doc, 'git clone https://github.com/rohitguptamyewards-lab/hyperlocal.git\ncd hyperlocal')

doc.add_paragraph()
add_para(doc, 'Step 2 — Backend setup', bold=True, size=11, color=(0x4F, 0x46, 0xE5))
add_code_block(doc,
    'cd src/backend\n'
    'composer install\n'
    '\n'
    '# .env is already committed (no need to copy .env.example)\n'
    '# SQLite database with all data is at: database/database.sqlite\n'
    '\n'
    'php artisan config:clear\n'
    'php artisan cache:clear\n'
    '\n'
    '# Start API server\n'
    'php artisan serve --port=8000'
)

doc.add_paragraph()
add_para(doc, 'Step 3 — Frontend setup', bold=True, size=11, color=(0x4F, 0x46, 0xE5))
add_code_block(doc,
    'cd src/frontend\n'
    'npm install\n'
    'npm run dev       # starts on http://localhost:3000'
)

doc.add_paragraph()
add_para(doc, 'Step 4 — Open the app', bold=True, size=11, color=(0x4F, 0x46, 0xE5))
add_table_with_headers(doc,
    ['URL', 'What You See'],
    [
        ['http://localhost:3000', 'Main merchant dashboard (login page first)'],
        ['http://localhost:8000/api/health', 'Backend health check (returns {"status":"ok"})'],
    ],
    col_widths=[2.5, 4.0]
)

doc.add_paragraph()
add_para(doc, 'Common Troubleshooting', bold=True, size=12, color=(0x4F, 0x46, 0xE5))

add_para(doc, 'Problem: "php not found" on Windows', bold=True, size=11)
add_bullet(doc, 'Add PHP to your system PATH: System Properties → Environment Variables → Path → Add PHP folder')

add_para(doc, 'Problem: "composer not found"', bold=True, size=11)
add_bullet(doc, 'Download from getcomposer.org and run the installer — it adds composer to PATH automatically')

add_para(doc, 'Problem: CORS error in browser', bold=True, size=11)
add_bullet(doc, 'Make sure the backend server is running on port 8000 and frontend on port 3000')
add_bullet(doc, 'Check src/backend/config/cors.php — allowed_origins should include http://localhost:3000')

add_para(doc, 'Problem: "Token mismatch" or auth errors after cloning', bold=True, size=11)
add_bullet(doc, 'Run: php artisan config:clear && php artisan cache:clear in src/backend')

doc.add_paragraph()
add_para(doc, 'Making Schema Changes', bold=True, size=12, color=(0x4F, 0x46, 0xE5))
add_numbered(doc, 'Create migration: php artisan make:migration add_column_to_table')
add_numbered(doc, 'Write the migration in database/migrations/')
add_numbered(doc, 'Run: php artisan migrate')
add_numbered(doc, 'Update the Model $fillable array')
add_numbered(doc, 'Update the Resource if you want the column in API responses')
add_numbered(doc, 'Update the TypeScript interface in src/frontend/src/stores/ or the relevant view')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SECTION 8 — GLOSSARY
# ════════════════════════════════════════════════════════

add_heading(doc, '8. Glossary', level=1, color=(0x4F, 0x46, 0xE5))
add_divider(doc)

add_table_with_headers(doc,
    ['Term', 'Definition'],
    [
        ['Partnership', 'A formal agreement between two brands to reward each other\'s customers'],
        ['Proposer', 'The brand that initiates/creates the partnership request'],
        ['Acceptor', 'The brand that receives the request and decides to accept or reject'],
        ['Offer Structure', '"Same" = one shared offer; "Different" = each brand configures their own'],
        ['Scope Type', '"Brand-wide" = applies to all outlets; "Outlet-level" = specific branches only'],
        ['POS Type', 'Point of Sale type: "flat" = fixed ₹ discount; "percentage" = % of bill'],
        ['Cap per Bill', 'Maximum ₹ discount allowed in a single transaction'],
        ['Monthly Cap', 'Maximum total ₹ discount a brand will give across all transactions in a month'],
        ['Min Bill', 'Minimum transaction amount for the offer to be applicable'],
        ['Offer Filled', 'Boolean flag — has a participant configured their offer yet? (for Different mode)'],
        ['Approval Mode', '1 = cashier auto-approves; 2 = manager must manually approve each redemption'],
        ['NEGOTIATING', 'Partnership status when either party has proposed updated terms'],
        ['Health Score', 'A metric (0–100) showing how actively a partnership is being used'],
        ['Network', 'A named group of merchants who can cross-promote to each other\'s customers'],
        ['Token', 'A unique code a customer presents at the POS to receive a partner discount'],
        ['Ledger', 'Double-entry accounting record of all point credits and debits'],
        ['Sanctum', 'Laravel\'s API token authentication system used throughout this platform'],
        ['Pinia', 'Vue 3\'s state management library (equivalent to Vuex/Redux)'],
        ['SQLite', 'Serverless SQL database stored as a single file — used for development'],
        ['UUID', 'Universally Unique ID — used in API URLs instead of integer IDs for security'],
    ],
    col_widths=[1.8, 4.7]
)

doc.add_paragraph()
doc.add_paragraph()

footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(
    f'Hyperlocal Network  •  Documentation v1.0  •  Generated {datetime.date.today().strftime("%d %B %Y")}'
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
footer_run.font.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────

output_path = r'C:\Users\rohit\Desktop\hyperlocal\Hyperlocal_Network_Documentation.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
